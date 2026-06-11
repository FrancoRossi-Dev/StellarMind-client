"""
export_code.py
Drop this file in the root of any C# solution and run:

    pip install pygments python-docx
    python export_code.py

Output: code_export.docx  (same folder as this script)
"""

import os
from pathlib import Path

from pygments import lex
from pygments.lexers import get_lexer_for_filename, TextLexer
from pygments.styles import get_style_by_name
from pygments.util import ClassNotFound

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── tweak these if needed ─────────────────────────────────────────────────────
EXTENSIONS    = {".cs", ".csproj", ".slnx", ".sln", ".json", ".cshtml"}
SKIP_DIRS     = {"bin", "obj", "Migrations", "Documentation", ".git", ".vs", ".claude"}
SKIP_SUFFIXES = {".Designer.cs"}   # EF / WinForms auto-generated files
SKIP_FILES    = set()
STYLE_NAME    = "vs"               # pygments style: "vs", "monokai", "friendly" …
FONT_NAME     = "Courier New"
FONT_SIZE     = Pt(8)
OUTPUT        = "code_export.docx"
# ─────────────────────────────────────────────────────────────────────────────


def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    if len(h) != 6:
        return RGBColor(0, 0, 0)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_code_file(doc: Document, filepath: Path, root: Path) -> None:
    style = get_style_by_name(STYLE_NAME)
    rel = filepath.relative_to(root)

    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run(str(rel))

    try:
        source = filepath.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        doc.add_paragraph(f"[Could not read file: {e}]")
        return

    try:
        lexer = get_lexer_for_filename(str(filepath))
    except ClassNotFound:
        lexer = TextLexer()

    tokens = list(lex(source, lexer))

    token_colors: dict = {}
    for ttype, color_info in style:
        if color_info.get("color"):
            token_colors[ttype] = color_info["color"]

    def get_color(ttype):
        t = ttype
        while t:
            if t in token_colors:
                return token_colors[t]
            t = t.parent if hasattr(t, "parent") else None
        return None

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    for ttype, value in tokens:
        if "\n" in value:
            parts = value.split("\n")
            for i, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE
                    color = get_color(ttype)
                    if color:
                        run.font.color.rgb = _rgb(color)
                if i < len(parts) - 1:
                    para = doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
        else:
            if value:
                run = para.add_run(value)
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                color = get_color(ttype)
                if color:
                    run.font.color.rgb = _rgb(color)


def collect_files(root: Path) -> list[Path]:
    files = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS]
        for fname in sorted(filenames):
            p = Path(dirpath) / fname
            if p.suffix in EXTENSIONS and not any(p.name.endswith(s) for s in SKIP_SUFFIXES) and p.name not in SKIP_FILES:
                files.append(p)
    return sorted(files)


def main() -> None:
    root = Path(__file__).parent.resolve()
    files = collect_files(root)

    if not files:
        print("No matching files found.")
        return

    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    for i, filepath in enumerate(files):
        print(f"  adding {filepath.relative_to(root)}")
        add_code_file(doc, filepath, root)
        if i < len(files) - 1:
            add_page_break(doc)

    doc.save(OUTPUT)
    print(f"Done -> {OUTPUT}")


if __name__ == "__main__":
    main()
